"""
Program Fit Evaluator — AI-assisted first-pass scoring of SB1 TCEP/SCCP grant nominations.

PMP-facing internal analyst tool. Ingests a nomination *package* (a 13-sheet "Program Fit"
Excel workbook + a Cal-B/C benefit-cost .xlsm + traffic studies + performance-metric backup
docs), scores each rubric factor 1-5 with a citation + confidence, screens eligibility and
cycle-deadline rules, flags contradictions and vague language, and exports a reviewable Excel
package with a human-in-the-loop review flow.

Architecturally this mirrors src/project_delivery_evaluator.py (PDE). It REUSES that module's
LLM client switcher (_get_client) and JSON salvage helper (_extract_json) by import, so the two
use cases stay consistent and nothing in PDE is modified.

See PROGRAM_FIT_EVALUATOR_SPEC.md (repo root) for the source-of-truth rubric, rules, and Cal-B/C
read strategy this module implements. OCR/vision fallback for scanned traffic tables is DEFERRED
(out of v1 scope).
"""

import datetime
import json
import logging
import os
import re
from io import BytesIO

import openpyxl

# Reuse PDE's battle-tested helpers verbatim — do NOT reimplement or modify them.
from src.project_delivery_evaluator import _get_client, _extract_json

logger = logging.getLogger(__name__)

# ==============================================================================
# RUBRIC  (spec §4 — recovered from PRC Results scoring sheet)
# ==============================================================================
# 10 scored criteria, each rated 1-5 (integer). Two groups: Freight System (3) and
# Transportation System (7). An 11th Air Quality / Community Impact criterion exists only in
# the feedback-loop tab and is intentionally kept OUT of the core 10 (spec §4.1).

# The 1-5 anchor definitions (spec §4.2). Applied to every criterion.
RATING_ANCHORS = {
    5: "High — factor is a PRIMARY objective; clear, direct, data-driven, significant benefits; "
       "quantified/data-backed.",
    4: "High (secondary) — strong benefits but on a secondary or indirect axis.",
    3: "Medium — not a primary purpose, or benefits don't meet a High bar; conceptual/qualitative, "
       "no supporting data.",
    2: "Low — insufficient information to assess benefits.",
    1: "Non-responsive — negatively impacts, or does not address the factor at all.",
}

# Freight System group (3 criteria)
_FREIGHT_CRITERIA = [
    {"id": "F1", "group": "Freight System", "name": "Throughput"},
    {"id": "F2", "group": "Freight System", "name": "Velocity"},
    {"id": "F3", "group": "Freight System", "name": "Reliability"},
]

# Transportation System group (7 criteria)
_TRANSPORT_CRITERIA = [
    {"id": "T1", "group": "Transportation System", "name": "Freight Safety"},
    {"id": "T2", "group": "Transportation System", "name": "Congestion Reduction/Mitigation"},
    {"id": "T3", "group": "Transportation System", "name": "Key Transportation Bottleneck Relief"},
    {"id": "T4", "group": "Transportation System", "name": "Multi-Modal Strategy"},
    {"id": "T5", "group": "Transportation System", "name": "Interregional Benefits"},
    {"id": "T6", "group": "Transportation System", "name": "Advanced Technology"},
    {"id": "T7", "group": "Transportation System", "name": "Freight Zero-Emission Infrastructure"},
]

# The core 10-criteria rubric. TCEP and SCCP share the same 10 criteria in v1; program-specific
# narrative emphasis differs but scoring structure is identical (spec §4).
RUBRIC_TCEP = _FREIGHT_CRITERIA + _TRANSPORT_CRITERIA
RUBRIC_SCCP = _FREIGHT_CRITERIA + _TRANSPORT_CRITERIA

# Overall score = AVERAGE of the 10 criteria, bucketed by rounding to nearest integer (spec §4.3).
# Each tuple: (inclusive_low, exclusive_high, label). Exact .5 boundaries yield split labels
# (handled in compute_program_fit_rating).
RATING_BUCKETS = [
    (4.5, 5.01, "HIGH"),
    (3.5, 4.5, "MEDIUM-HIGH"),
    (2.5, 3.5, "MEDIUM"),
    (1.5, 2.5, "MEDIUM-LOW"),
    (0.0, 1.5, "LOW"),
]

# Per-criterion calibration averages from the PRC population (spec §4.5) — reference only.
# These are the ACTUAL before-feedback PRC averages ('PRC Results (before feedback)' row 40).
# NOTE: the 'Potential Results-feedback loop' tab row 39 carries much higher hypothetical
# "if responsive" averages (e.g. Safety 4.80, ZEV 3.90) — do not confuse the two.
CALIBRATION_AVERAGES = {
    "Throughput": 2.87, "Velocity": 2.97, "Reliability": 2.83, "Freight Safety": 3.27,
    "Congestion Reduction/Mitigation": 2.97, "Key Transportation Bottleneck Relief": 2.97,
    "Multi-Modal Strategy": 2.83, "Interregional Benefits": 3.20, "Advanced Technology": 2.03,
    "Freight Zero-Emission Infrastructure": 1.13,
}

# ==============================================================================
# ELIGIBILITY / CYCLE RULES  (spec §5 — deterministic, codeable)
# ==============================================================================
# Ineligibility conditions, verbatim from Risk Ratings B9 (spec §5.4). (*) = SCCP-only.
INELIGIBILITY_CONDITIONS = [
    "Uncommitted funds must be only from: SCCP, TCEP, LPP, or a federal discretionary grant program.",
    "Construction phase of a capital project only. (*SCCP)",
    "Environmental process must be complete within 6 months after program adoption.",
    "Must not supplant other committed funds.",
    "Fund cost increases (SCCP only). (*SCCP)",
    "Implementing agency must prove ability to absorb overruns with no additional SCCP/TCEP funding.",
    "Project must have independent utility / standalone corridor benefits.",
    "Project must be ready to start construction by December 31, 2029.",
    "Caltrans does not pay for SCCP cost overruns. (*SCCP)",
]

# Cycle deadlines (spec §5.1). NOTE: Appendix A says SCCP construction-ready by Dec 31 2027 while
# the main form says Dec 31 2029 — unresolved SOURCE contradiction flagged for business (spec §10).
CYCLE_DEADLINES = {
    "paed_complete_by": "2027-12-31",       # environmental clearance, within 6mo of program adoption
    "rtl_by": "2029-06-30",                 # ready-to-list for CON-funded projects
    "construction_ready_by": "2029-12-31",  # main form governs; Appendix A conflict noted
}

# Funding risk tiers (spec §5.3). Dollar figures in whole dollars.
FUNDING_RISK_TIERS = {
    "High": "uncommitted > $60M, OR request > $90M, OR no leverage/match.",
    "Medium": "$25-60M uncommitted, OR $75-89M request, OR leverage < 10%.",
    "Low": "< $25M uncommitted, OR request <= $75M.",
}
TCEP_MIN_MATCH_RATE = 0.30   # TCEP requires 30% minimum local match on the regional request (spec §5.3)

# Schedule risk tiers (spec §5.2) — RTL / CON windows.
SCHEDULE_RISK_TIERS = {
    "High": "RTL Jan-Jun 2029 / CON Jun-Dec 2029 (closest to deadline).",
    "Medium": "RTL Jul-Dec 2028 / CON Dec 2028-May 2029.",
    "Low": "RTL before Jul 2027-Jun 2028 / CON Jul 2027-Nov 2028.",
}

# Cal-B/C output defined names — identical meaning across Sketch and Corridor editions (spec §6.2).
CALBC_DEFINED_NAMES = [
    "LifeCycleCost", "LifeCycleBene", "NetPresentValue", "ReturnOnInvest", "Payback", "DiscRate",
]


# ==============================================================================
# PACKAGE EXTRACTION  (readers)
# ==============================================================================
def extract_package(files) -> dict:
    """Ingest a nomination package (mixed file types) into a structured dict.

    Dispatches by file extension: .xlsx -> read_program_fit_workbook, .xlsm -> read_calbc_model,
    .pdf/.docx -> narrative + table text (clean-text only; scanned-appendix OCR is deferred).

    Args:
        files: list of uploaded file-like objects (Streamlit UploadedFile) or paths.

    Returns:
        {"workbook": dict|None, "calbc": dict|None, "narratives": [ {name, text} ],
         "metric_tables": [...], "warnings": [...], "retriever": HybridRetriever|None,
         "ocr_sources": [names]}
    """
    from src.project_delivery_evaluator import (
        extract_text_from_docx,
        extract_text_from_uploaded_pdf,
    )

    result = {
        "workbook": None,
        "calbc": None,
        "narratives": [],
        "metric_tables": [],
        "warnings": [],
        "retriever": None,
        "ocr_sources": [],
    }

    # Text recovered by OCR from scanned PDFs (best-effort; feeds the retriever index).
    ocr_corpus = []  # list of (name, full_text)

    def _name_of(f):
        n = getattr(f, "name", None)
        if n:
            return str(n)
        return str(f) if isinstance(f, str) else "unknown"

    def _rewind(f):
        # Reset stream position so a file-like object can be read again.
        try:
            f.seek(0)
        except (AttributeError, ValueError, OSError):
            pass
        return f

    for f in files or []:
        name = _name_of(f)
        base = os.path.basename(name)
        ext = os.path.splitext(name)[1].lower()
        try:
            if ext == ".xlsx":
                result["workbook"] = read_program_fit_workbook(_rewind(f))
            elif ext == ".xlsm":
                result["calbc"] = read_calbc_model(_rewind(f))
            elif ext == ".docx":
                tables = extract_docx_tables(_rewind(f))
                result["metric_tables"].append({"name": base, "tables": tables})
                text = extract_text_from_docx(_rewind(f))
                result["narratives"].append({"name": base, "text": text})
            elif ext == ".pdf":
                # Fast clean-text path first (unchanged behavior, controls OCR cost).
                text = extract_text_from_uploaded_pdf(_rewind(f))
                result["narratives"].append({"name": base, "text": text})
                # If the PDF has little/no extracted text it is likely a scanned appendix;
                # OCR it best-effort and collect the recovered text for retrieval. An OCR
                # failure (missing tesseract/PyMuPDF, etc.) must never sink the package.
                if len((text or "").strip()) < 100:
                    try:
                        from src.program_fit_ocr import extract_pdf_text
                        ocr_out = extract_pdf_text(
                            _rewind(f), cache_dir=os.getenv("PROGRAM_FIT_CACHE_DIR") or None
                        )
                        ocr_text = ocr_out.get("full_text", "") if isinstance(ocr_out, dict) else ""
                        if ocr_text and ocr_text.strip():
                            ocr_corpus.append((base, ocr_text))
                    except Exception as oe:  # noqa: BLE001 — OCR is best-effort
                        result["warnings"].append(f"{base}: OCR skipped: {type(oe).__name__}: {oe}")
            else:
                result["warnings"].append(f"{base}: unsupported file type '{ext}'")
        except Exception as e:  # noqa: BLE001 — one bad file must not sink the package
            result["warnings"].append(f"{base}: {type(e).__name__}: {e}")

    # ---- Build a retrieval index over any OCR'd traffic-study text ----
    if ocr_corpus:
        try:
            import hashlib
            import tempfile
            from src.program_fit_retrieval import build_or_load_index

            combined = "\n\n".join(t for _, t in ocr_corpus)
            key = hashlib.sha1(combined.encode("utf-8", "replace")).hexdigest()[:16]
            cache_dir = os.getenv("PROGRAM_FIT_CACHE_DIR") or os.path.join(
                tempfile.gettempdir(), "pf_retrieval_cache"
            )
            result["retriever"] = build_or_load_index(combined, cache_dir, key)
            result["ocr_sources"] = [nm for nm, _ in ocr_corpus]
        except Exception as e:  # noqa: BLE001 — retriever build is best-effort
            result["warnings"].append(f"retriever build skipped: {type(e).__name__}: {e}")

    return result


def read_program_fit_workbook(xlsx) -> dict:
    """Read the 13-sheet Program Fit .xlsx into structured form data.

    Reads the orange "SB1 Cycle 5 Program Fit" sheet (answers in col D, question text in col C),
    the Section V performance-measure tables (cols G-M), and Section IV funding figures.

    Returns:
        {"project_name": str, "district": str, "program": "TCEP"|"SCCP"|"BOTH",
         "answers": {line_id: text}, "funding": {...}, "performance_tables": {...},
         "form_id": str}
    """
    stream = xlsx
    if hasattr(xlsx, "read"):
        try:
            xlsx.seek(0)
        except (AttributeError, ValueError, OSError):
            pass
        stream = BytesIO(xlsx.read())
    wb = openpyxl.load_workbook(stream, data_only=True)

    # Locate the orange main sheet: contains "Program Fit" but is NOT the instructions tab.
    main_name = None
    for sn in wb.sheetnames:
        if "program fit" in sn.lower() and "instruction" not in sn.lower():
            main_name = sn
            break
    if main_name is None:
        for sn in wb.sheetnames:
            if "program fit" in sn.lower():
                main_name = sn
                break
    ws = wb[main_name] if main_name else wb[wb.sheetnames[0]]

    def _cell(coord):
        try:
            return ws[coord].value
        except Exception:
            return None

    def _s(v):
        return str(v).strip() if v is not None else ""

    project_name = _s(_cell("D15"))
    district_raw = _cell("D17")
    district = _s(district_raw)
    form_id = _s(_cell("A14"))

    tcep_flag = bool(_cell("D9"))
    sccp_flag = bool(_cell("D11"))

    # Section IV funding figures (values in $1,000s) with their source cells.
    def _figure(coord):
        return {"value": _cell(coord), "cell": f"'{main_name}'!{coord}"}

    tcep_request = _figure("D45")
    tcep_regional = _figure("D46")
    tcep_state = _figure("D48")
    sccp_request = _figure("D43")
    uncommitted = _figure("D52")

    # Program inference: prefer the checkbox flags, fall back to Section IV request figures.
    if tcep_flag and sccp_flag:
        program = "BOTH"
    elif tcep_flag:
        program = "TCEP"
    elif sccp_flag:
        program = "SCCP"
    else:
        has_tcep = isinstance(tcep_request["value"], (int, float)) and tcep_request["value"]
        has_sccp = isinstance(sccp_request["value"], (int, float)) and sccp_request["value"]
        if has_tcep and has_sccp:
            program = "BOTH"
        elif has_sccp:
            program = "SCCP"
        else:
            program = "TCEP"

    # Fund-table (Section IV backup grid) lives on the "Project Funding Info #1" sheet:
    # 8 blocks, label at AA{r}, total at AH{r+9}.
    fund_name = None
    for sn in wb.sheetnames:
        if "project funding info" in sn.lower():
            fund_name = sn
            break
    fund_ws = wb[fund_name] if fund_name else ws
    fund_label = fund_name or main_name

    def _fcell(coord):
        try:
            return fund_ws[coord].value
        except Exception:
            return None

    fund_table_blocks = []
    tcep_sccp_total = 0
    tcep_sccp_cells = []
    for label_row in (18, 29, 40, 51, 62, 73, 84, 95):
        total_row = label_row + 9
        label = _s(_fcell(f"AA{label_row}"))
        total = _fcell(f"AH{total_row}")
        block = {
            "label": label,
            "label_cell": f"'{fund_label}'!AA{label_row}",
            "total": total,
            "total_cell": f"'{fund_label}'!AH{total_row}",
        }
        fund_table_blocks.append(block)
        lu = label.upper()
        if ("TCEP" in lu or "SCCP" in lu) and isinstance(total, (int, float)):
            tcep_sccp_total += total
            tcep_sccp_cells.append(block["total_cell"])

    funding = {
        "tcep_request": tcep_request,
        "tcep_regional_request": tcep_regional,
        "tcep_state_request": tcep_state,
        "sccp_request": sccp_request,
        "uncommitted": uncommitted,
        "fund_table_blocks": fund_table_blocks,
        "fund_table_tcep_sccp_total": tcep_sccp_total,
        "fund_table_tcep_sccp_total_cells": tcep_sccp_cells,
    }

    # Date cells on the funding-info sheet (PPR milestone schedule). The impossible 1930s
    # closeout dates in the real samples live HERE (e.g. 'Project Funding Info #1'!V33), not on
    # the main form, so detect_contradictions needs them surfaced.
    funding_milestones = []
    if fund_ws is not ws:
        for frow in fund_ws.iter_rows():
            for cell in frow:
                v = cell.value
                if isinstance(v, (datetime.datetime, datetime.date)):
                    funding_milestones.append({
                        "cell": f"'{fund_label}'!{cell.coordinate}",
                        "value": v.isoformat()[:10],
                        "year": v.year,
                    })

    # Answers: question (col C) -> answer (col D), keyed by row, for rows that have content.
    answers = {}
    for row in range(1, ws.max_row + 1):
        q = _cell(f"C{row}")
        a = _cell(f"D{row}")
        if a is None:
            continue
        qs = _s(q)
        if not qs and not _s(a):
            continue
        answers[str(row)] = {"question": qs, "answer": _cell(f"D{row}")}

    return {
        "project_name": project_name,
        "district": district,
        "program": program,
        "answers": answers,
        "funding": funding,
        "funding_milestones": funding_milestones,
        "performance_tables": {},
        "form_id": form_id,
    }


def read_calbc_model(xlsm) -> dict:
    """Read a Cal-B/C benefit-cost .xlsm via workbook-defined NAMES (not fixed cells).

    Loads with openpyxl(data_only=True, keep_vba=True), detects edition from the Title sheet
    ("Sketch" vs "Corridor"), and resolves CALBC_DEFINED_NAMES to cached values. Branches
    itemized/emissions grids on edition (spec §6).

    Returns:
        {"edition": "Sketch"|"Corridor", "benefit_cost_ratio": float, "life_cycle_cost": float,
         "life_cycle_benefit": float, "net_present_value": float, "payback": str|float,
         "discount_rate": float, "raw_named_values": {name: value}}
    """
    stream = xlsm
    if hasattr(xlsm, "read"):
        try:
            xlsm.seek(0)
        except (AttributeError, ValueError, OSError):
            pass
        stream = BytesIO(xlsm.read())
    wb = openpyxl.load_workbook(stream, data_only=True, keep_vba=True)

    # --- Edition detection: scan the Title sheet (spec: Title!C2) for Sketch/Corridor ---
    edition = None
    title_ws = None
    for sn in wb.sheetnames:
        if sn.strip().lower() == "title":
            title_ws = wb[sn]
            break
    if title_ws is None:
        title_ws = wb[wb.sheetnames[0]]

    def _classify(text):
        t = (text or "").lower()
        if "sketch" in t:
            return "Sketch"
        if "corridor" in t:
            return "Corridor"
        return None

    try:
        edition = _classify(str(title_ws["C2"].value))
    except Exception:
        edition = None
    if edition is None:
        # Fall back to scanning the whole Title sheet.
        for row in title_ws.iter_rows(values_only=True):
            for val in row:
                if isinstance(val, str):
                    ed = _classify(val)
                    if ed:
                        edition = ed
                        break
            if edition:
                break

    # --- Resolve a defined NAME to its cached (data_only) value ---
    def _resolve_named(name):
        dn = None
        try:
            if hasattr(wb.defined_names, "get"):
                dn = wb.defined_names.get(name)
            else:
                dn = wb.defined_names[name]
        except Exception:
            dn = None
        if dn is None:
            return None
        try:
            for sheet, coord in dn.destinations:
                try:
                    return wb[sheet][coord].value
                except Exception:
                    continue
        except Exception:
            return None
        return None

    raw_named_values = {n: _resolve_named(n) for n in CALBC_DEFINED_NAMES}

    # --- Benefit/Cost ratio: prefer the BeneCostRatio defined name, else read '3) Results' ---
    bc_ratio = _resolve_named("BeneCostRatio")
    if bc_ratio is None:
        results_ws = None
        for sn in wb.sheetnames:
            if "results" in sn.lower():
                results_ws = wb[sn]
                break
        if results_ws is not None:
            for row in results_ws.iter_rows():
                for cell in row:
                    val = cell.value
                    if isinstance(val, str) and "benefit" in val.lower() and "cost ratio" in val.lower():
                        cand = results_ws.cell(row=cell.row, column=8).value  # column H
                        if isinstance(cand, (int, float)):
                            bc_ratio = cand
                        break
                if bc_ratio is not None:
                    break

    def _num(v):
        return v if isinstance(v, (int, float)) else None

    return {
        "edition": edition,
        "benefit_cost_ratio": _num(bc_ratio),
        "life_cycle_cost": _num(raw_named_values.get("LifeCycleCost")),
        "life_cycle_benefit": _num(raw_named_values.get("LifeCycleBene")),
        "net_present_value": _num(raw_named_values.get("NetPresentValue")),
        "payback": raw_named_values.get("Payback"),
        "discount_rate": _num(raw_named_values.get("DiscRate")),
        "raw_named_values": raw_named_values,
    }


def extract_docx_tables(docx) -> list:
    """Extract tables (not just paragraphs) from a performance-metric backup .docx.

    The existing extract_text_from_docx reads paragraphs only; this adds doc.tables so metric ->
    BCA-cell traceability rows are captured (spec §6.3).

    Returns: list of tables, each a list of row-lists of cell strings.
    """
    from docx import Document

    if hasattr(docx, "read"):
        try:
            docx.seek(0)
        except (AttributeError, ValueError, OSError):
            pass
        doc = Document(BytesIO(docx.read()))
    else:
        doc = Document(docx)

    return [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]


# ==============================================================================
# LLM SCORING  (single call, mirrors PDE contract)
# ==============================================================================
def run_program_fit_evaluation(package: dict, program: str = "TCEP", model_name: str = "gpt-4o",
                               pf_rules=None) -> dict:
    """Score all 10 rubric criteria in ONE LLM call, each with citation + confidence.

    Mirrors PDE's run_delivery_evaluation contract: builds a system prompt with the rubric +
    anchors, sends the extracted package text, parses via _extract_json with one retry on
    malformed JSON. Confidence < 0.5 forces the criterion to "insufficient information" (rating 2)
    and missing_info=True.

    Returns:
        {"project_name": str, "district": str, "program": str, "evaluation_date": "YYYY-MM-DD",
         "criteria": [ {criterion_id, name, group, score(1-5), source_reasoning, missing_info,
                        confidence(0-1)} ],  # exactly 10
         "missing_criteria": [id...], "summary": str}  or {"error": "..."}
    """
    try:
        program = (program or "TCEP").upper()

        # ---- Institutional memory (prior reviewer calibrations) injected into the prompt ----
        try:
            from src.program_fit_memory_manager import build_institutional_memory_block
            memory_block = build_institutional_memory_block(pf_rules or [])
        except Exception:  # noqa: BLE001 — memory must never break scoring
            memory_block = ""

        system_prompt = _build_system_prompt(program, memory_block=memory_block)

        # ---- Assemble the user message from the package (bounded to ~40k chars) ----
        package = package or {}
        MAX_TOTAL = 40_000
        parts = []
        used = 0

        def _add(header, text):
            nonlocal used
            if used >= MAX_TOTAL or not text:
                return
            text = str(text)
            budget = MAX_TOTAL - used
            chunk = text[:budget]
            block = f"{header}\n{chunk}"
            parts.append(block)
            used += len(block)

        # ---- Retrieval: pull rubric-relevant excerpts from OCR'd traffic studies ----
        retriever = package.get("retriever")
        if retriever is not None:
            try:
                rubric = RUBRIC_TCEP if program == "TCEP" else RUBRIC_SCCP
                query = " ".join(c["name"] for c in rubric) + (
                    " freight congestion delay level of service throughput safety"
                )
                hits = retriever.retrieve(query, k=8) or []
                joined = "\n\n".join(
                    str(h.get("chunk", "")) for h in hits if isinstance(h, dict)
                )
                _add("=== RETRIEVED TRAFFIC-STUDY EXCERPTS ===", joined)
            except Exception:  # noqa: BLE001 — a retriever error must never break scoring
                pass

        # Workbook answers
        workbook = package.get("workbook")
        if isinstance(workbook, dict):
            answers = workbook.get("answers")
            if isinstance(answers, dict):
                lines = []
                for key, entry in answers.items():
                    if isinstance(entry, dict):
                        q = str(entry.get("question", "")).strip()
                        a = entry.get("answer")
                        lines.append(f"[{key}] {q}: {a}")
                    else:
                        lines.append(f"[{key}] {entry}")
                _add("=== PROGRAM FIT WORKBOOK ANSWERS ===", "\n".join(lines))
        elif workbook:
            _add("=== PROGRAM FIT WORKBOOK ANSWERS ===", workbook)

        # Narratives (list of {name, text})
        narratives = package.get("narratives")
        if isinstance(narratives, (list, tuple)):
            for nar in narratives:
                if isinstance(nar, dict):
                    _add(f"=== NARRATIVE: {nar.get('name', 'narrative')} ===", nar.get("text", ""))
                elif nar:
                    _add("=== NARRATIVE ===", nar)
        elif isinstance(narratives, dict):
            _add(f"=== NARRATIVE: {narratives.get('name', 'narrative')} ===", narratives.get("text", ""))
        elif narratives:
            _add("=== NARRATIVE ===", narratives)

        # Performance-metric backup tables (.docx). The real backup docs are often table-only —
        # extract_text_from_docx returns "" for them — so the tables must be rendered here or the
        # submitted metrics are invisible to scoring.
        metric_tables = package.get("metric_tables")
        if isinstance(metric_tables, (list, tuple)):
            for mt in metric_tables:
                if not isinstance(mt, dict):
                    continue
                rendered_tables = []
                for table in (mt.get("tables") or []):
                    try:
                        rows = [" | ".join(str(c) for c in row) for row in table]
                    except TypeError:
                        continue
                    if rows:
                        rendered_tables.append("\n".join(rows))
                if rendered_tables:
                    _add(
                        f"=== PERFORMANCE-METRIC BACKUP TABLES: {mt.get('name', 'backup doc')} ===",
                        "\n\n".join(rendered_tables),
                    )

        # Cal-B/C benefit-cost summary
        calbc = package.get("calbc")
        if calbc:
            _add("=== CAL-B/C BENEFIT-COST SUMMARY ===", json.dumps(calbc, default=str))

        package_text = "\n\n".join(parts) if parts else "(no package content provided)"

        user_message = (
            f"Please evaluate the following {program} nomination package against all 10 rubric "
            f"criteria. Score each criterion 1-5 with an exact-quote citation, missing_info flag, "
            f"and confidence.\n\nNOMINATION PACKAGE CONTENT:\n{package_text}"
        )

        response_format = (
            {"type": "json_object"}
            if any(m in model_name.lower() for m in ["gpt", "json", "gemma"])
            else None
        )

        client = _get_client(model_name)
        response = client.chat.completions.create(
            model=model_name,
            response_format=response_format,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.0,
        )

        main_text = response.choices[0].message.content
        main_reason = response.choices[0].finish_reason

        try:
            return _extract_json(main_text, finish_reason=main_reason)
        except (json.JSONDecodeError, ValueError):
            # Retry once on malformed JSON
            retry = client.chat.completions.create(
                model=model_name,
                response_format=response_format,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message},
                ],
                temperature=0.0,
            )
            retry_text = retry.choices[0].message.content
            retry_reason = retry.choices[0].finish_reason
            return _extract_json(retry_text, finish_reason=retry_reason)
    except Exception as e:  # noqa: BLE001 — must degrade to an error dict, never raise
        return {"error": f"AI service error during program fit evaluation: {str(e)}"}


def _build_system_prompt(program: str, kb_context: str = "", memory_block: str = "") -> str:
    """Assemble the scoring system prompt: persona + rubric + 1-5 anchors + confidence rules +
    output JSON schema. Mirrors PDE _build_system_prompt structure.

    `memory_block` (optional) is a pre-rendered institutional-memory block of prior reviewer
    calibrations (see program_fit_memory_manager.build_institutional_memory_block). When
    non-empty it is appended after the rubric so past calibrations guide scoring; when "" the
    prompt is byte-for-byte identical to the pre-HIFL behavior."""
    program = (program or "TCEP").upper()
    rubric = RUBRIC_TCEP if program == "TCEP" else RUBRIC_SCCP

    persona = f"""You are a Caltrans Programming and Regional Coordination (PRC) first-pass reviewer for SB1 {program} grant nominations. You have deep experience scoring competitive freight and corridor project nominations against the official {program} rubric.

You are meticulous, evidence-based, and transparent about uncertainty. When the nomination package lacks information for a criterion, you flag it clearly rather than guessing confidently."""

    # --- Embed the rubric verbatim (every criterion name + group must appear) ---
    rubric_lines = [f"{program} SCORING RUBRIC — 10 criteria across two groups. Score EACH criterion 1-5:"]
    current_group = None
    for c in rubric:
        if c["group"] != current_group:
            current_group = c["group"]
            rubric_lines.append(f"\nGROUP: {current_group}")
        rubric_lines.append(f'  {c["id"]} — {c["name"]}  (group: {c["group"]})')
    rubric_text = "\n".join(rubric_lines)

    # --- 1-5 anchor guidance (must contain digits 1..5 and the word "primary") ---
    anchor_lines = ["RATING ANCHORS (apply the 1-5 scale to every criterion):"]
    for level in (5, 4, 3, 2, 1):
        anchor_lines.append(f"  {level} = {RATING_ANCHORS[level]}")
    anchors_text = "\n".join(anchor_lines)

    confidence_rule = """CONFIDENCE & MISSING-INFORMATION RULE (MANDATORY):
- confidence is a float in [0,1] reflecting how certain you are about the CORRECT score for the criterion, based strictly on the directness of the quoted evidence.
- If confidence < 0.5, set the score to 2 (insufficient information) and missing_info=true. A confidence below 0.5 means you are more uncertain than certain, so the honest score is 2 (Low — insufficient information to assess benefits).
- If the package contains no direct evidence for a criterion, set source_reasoning to exactly "No direct evidence found", confidence to a low value (< 0.25), score to 2, and missing_info=true.
- Do NOT inflate confidence to avoid this gate."""

    output_schema = """OUTPUT FORMAT — return ONLY a single JSON object with EXACTLY this schema:
{
  "project_name": "<project name>",
  "district": "<district>",
  "program": "<TCEP or SCCP>",
  "evaluation_date": "YYYY-MM-DD",
  "criteria": [
    {
      "criterion_id": "<e.g. F1>",
      "name": "<criterion name>",
      "group": "<Freight System or Transportation System>",
      "score": <integer 1-5>,
      "source_reasoning": "<quote the exact document text supporting the score, or \\"No direct evidence found\\">",
      "missing_info": <true or false>,
      "confidence": <float 0-1>
    }
    // ... EXACTLY 10 criteria objects, in the rubric order shown above
  ],
  "missing_criteria": ["<criterion_id of each criterion with missing_info=true>"],
  "summary": "<a concise overall program-fit summary>"
}
Do not include any text outside the JSON object."""

    sections = [persona, rubric_text, anchors_text, confidence_rule, output_schema]
    if kb_context:
        sections.insert(1, f"REFERENCE KNOWLEDGE BASE:\n{kb_context}")
    if memory_block:
        # Insert prior reviewer calibrations right after the rubric so they clearly scope how
        # each criterion is scored, without overriding the confidence/output-schema rules.
        rubric_idx = sections.index(rubric_text)
        sections.insert(rubric_idx + 1, memory_block)

    return "\n\n".join(sections)


# ==============================================================================
# SCORING / RATING  (pure Python — deterministic, unit-testable)
# ==============================================================================
def score_all_factors(evaluation: dict) -> dict:
    """Normalize/validate the per-criterion LLM scores into a clean scores dict.

    Any of the 10 canonical criteria the LLM omitted (or returned a non-numeric score for) is
    padded to 2 ("insufficient information", spec §4.2) so the overall average is always taken
    over all 10 criteria — a smaller denominator would silently inflate the rating.

    Returns: {"scores": {criterion_id: int 1-5} (all 10), "by_group": {...},
              "count": <criteria actually scored by the LLM>, "missing": [ids padded to 2]}
    """
    evaluation = evaluation or {}
    criteria = evaluation.get("criteria") or []

    # id -> group lookup from the canonical rubric
    id_to_group = {c["id"]: c["group"] for c in RUBRIC_TCEP}
    valid_ids = set(id_to_group)

    scores = {}
    by_group = {"Freight System": {}, "Transportation System": {}}
    count = 0

    for crit in criteria:
        if not isinstance(crit, dict):
            continue
        cid = crit.get("criterion_id", crit.get("id"))
        if cid is None or cid not in valid_ids:
            continue
        raw = crit.get("score")
        try:
            score = int(round(float(raw)))
        except (TypeError, ValueError):
            continue
        count += 1
        # clamp to 1-5
        score = max(1, min(5, score))
        scores[cid] = score
        by_group[id_to_group[cid]][cid] = score

    # Pad any of the 10 criteria the LLM did not usably score with 2 (insufficient information).
    missing = []
    for cid, group in id_to_group.items():
        if cid not in scores:
            scores[cid] = 2
            by_group[group][cid] = 2
            missing.append(cid)

    return {"scores": scores, "by_group": by_group, "count": count, "missing": missing}


def compute_program_fit_rating(scores) -> dict:
    """Overall = AVERAGE of the 10 criteria, bucketed via RATING_BUCKETS (spec §4.3).

    Args:
        scores: dict {criterion_id: 1-5} or list of 10 ints.

    Returns:
        {"average": float, "rating": str, "split": bool}
        (split=True with a combined label like "MEDIUM/MEDIUM-HIGH" on exact .5 boundaries.)
    """
    # Normalize input to a flat list of ints
    if isinstance(scores, dict):
        values = list(scores.values())
    else:
        values = list(scores)
    values = [int(v) for v in values]

    # Guard against a malformed evaluation that produced no usable criteria — never divide by
    # zero; surface an explicit INCOMPLETE rating the caller can handle.
    if not values:
        return {"average": 0.0, "rating": "INCOMPLETE", "split": False}

    average = sum(values) / len(values)

    def _bucket_for(avg):
        for low, high, label in RATING_BUCKETS:
            if low <= avg < high:
                return label
        return RATING_BUCKETS[0][2]  # fallback (avg at very top)

    # Exact .5 interior boundaries produce split labels (lower bucket first).
    # 4.5 is the top boundary into HIGH and is intentionally NOT split.
    split_boundaries = {3.5, 2.5, 1.5}
    if average in split_boundaries:
        lower_label = next(lbl for low, high, lbl in RATING_BUCKETS if high == average)
        higher_label = next(lbl for low, high, lbl in RATING_BUCKETS if low == average)
        return {
            "average": float(average),
            "rating": f"{lower_label}/{higher_label}",
            "split": True,
        }

    return {
        "average": float(average),
        "rating": _bucket_for(average),
        "split": False,
    }


# ==============================================================================
# RULES ENGINE  (deterministic checks)
# ==============================================================================
def _parse_answer_date(v):
    """Coerce a workbook answer to a date. Handles real date cells AND the string forms the
    filled samples actually use — D10's milestones are strings like '08/2029' — returning the
    first of the month when no day is given. Returns None when the value isn't a date."""
    if isinstance(v, datetime.datetime):
        return v.date()
    if isinstance(v, datetime.date):
        return v
    if not isinstance(v, str):
        return None
    s = v.strip()
    if not s:
        return None
    for fmt in ("%m/%d/%Y", "%m-%d-%Y", "%Y-%m-%d", "%Y/%m/%d"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    m = re.fullmatch(r"(\d{1,2})[/-](\d{4})", s)  # month/year, e.g. '08/2029'
    if m:
        month, year = int(m.group(1)), int(m.group(2))
        if 1 <= month <= 12:
            return datetime.date(year, month, 1)
    m = re.fullmatch(r"([A-Za-z]{3,9})\.?\s+(\d{4})", s)  # 'Aug 2029' / 'August 2029'
    if m:
        try:
            month = datetime.datetime.strptime(m.group(1)[:3].title(), "%b").month
        except ValueError:
            return None
        return datetime.date(int(m.group(2)), month, 1)
    return None


def screen_eligibility(package: dict) -> dict:
    """Apply deterministic ineligibility + cycle-deadline rules (spec §5).

    Returns:
        {"eligible": bool, "failures": [ {rule, detail} ], "schedule_risk": str|None,
         "funding_risk": str|None, "warnings": [...]}
    """
    result = {
        "eligible": True,
        "failures": [],
        "schedule_risk": None,
        "funding_risk": None,
        "warnings": [],
    }
    if not isinstance(package, dict):
        return result
    workbook = package.get("workbook")
    if not isinstance(workbook, dict):
        result["warnings"].append("No Program Fit workbook found in package; eligibility not screened.")
        return result

    funding = workbook.get("funding")
    if not isinstance(funding, dict):
        funding = {}
    answers = workbook.get("answers")
    if not isinstance(answers, dict):
        answers = {}
    program = str(workbook.get("program") or "").upper()

    def _num(v):
        return v if isinstance(v, (int, float)) and not isinstance(v, bool) else None

    def _fig(key):
        fig = funding.get(key)
        if isinstance(fig, dict):
            return _num(fig.get("value"))
        return _num(fig)

    # ---- Funding risk (values stored in $1,000s -> convert to whole dollars) ----
    tcep_req_k = _fig("tcep_request")
    sccp_req_k = _fig("sccp_request")
    uncommitted_k = _fig("uncommitted")
    request_candidates = [v for v in (tcep_req_k, sccp_req_k) if v is not None]
    request_k = max(request_candidates) if request_candidates else None
    request_dollars = request_k * 1000 if request_k is not None else None
    uncommitted_dollars = uncommitted_k * 1000 if uncommitted_k is not None else None

    def _classify_funding():
        if uncommitted_dollars is not None and uncommitted_dollars > 60_000_000:
            return "High"
        if request_dollars is not None and request_dollars > 90_000_000:
            return "High"
        if uncommitted_dollars is not None and uncommitted_dollars >= 25_000_000:
            return "Medium"
        if request_dollars is not None and request_dollars >= 75_000_000:
            return "Medium"
        if request_dollars is not None:
            return "Low"
        return None

    result["funding_risk"] = _classify_funding()

    # ---- Locate milestone dates in the answers (construction / RTL) ----
    def _find_date(keywords):
        for entry in answers.values():
            if not isinstance(entry, dict):
                continue
            q = str(entry.get("question", "")).lower()
            if any(kw in q for kw in keywords):
                d = _parse_answer_date(entry.get("answer"))
                if d is not None:
                    return d
        return None

    con_date = _find_date(["begin construction", "construction (month", "start construction"])
    rtl_date = _find_date(["rtl", "ready to list", "ready-to-list"])

    # ---- Schedule risk (spec §5.2 windows) ----
    def _classify_schedule():
        if con_date is not None:
            y, m = con_date.year, con_date.month
            if y == 2029 and 6 <= m <= 12:
                return "High"
            if (y == 2028 and m == 12) or (y == 2029 and 1 <= m <= 5):
                return "Medium"
            if con_date <= datetime.date(2028, 11, 30):
                return "Low"
            return None
        if rtl_date is not None:
            y, m = rtl_date.year, rtl_date.month
            if y == 2029 and 1 <= m <= 6:
                return "High"
            if y == 2028 and 7 <= m <= 12:
                return "Medium"
            if rtl_date <= datetime.date(2028, 6, 30):
                return "Low"
        return None

    result["schedule_risk"] = _classify_schedule()

    # ---- Hard rule: must be ready to start construction by 2029-12-31 (spec §5.4 / §5.1) ----
    deadline = datetime.date(2029, 12, 31)
    if con_date is not None and con_date > deadline:
        result["eligible"] = False
        result["failures"].append({
            "rule": "Project must be ready to start construction by December 31, 2029.",
            "detail": (
                f"Target begin-construction date {con_date.isoformat()} is after the "
                f"{deadline.isoformat()} cycle deadline."
            ),
        })

    # ---- TCEP minimum local-match note (soft, not a hard failure) ----
    if program in ("TCEP", "BOTH"):
        regional_k = _fig("tcep_regional_request")
        if regional_k is not None:
            result["warnings"].append(
                f"TCEP nomination: verify the regional request carries at least a "
                f"{int(TCEP_MIN_MATCH_RATE * 100)}% local match (spec §5.3); match not confirmable "
                f"from the workbook figures alone."
            )

    return result


def detect_contradictions(package: dict) -> list:
    """Cross-field contradiction checks (spec §8). E.g. Section IV TCEP request vs fund-table total
    (D10: $65M vs $60M), units anomalies, impossible milestone dates.

    Returns: list of {"field", "detail", "values", "severity"}.
    """
    contradictions = []
    if not isinstance(package, dict):
        return contradictions
    workbook = package.get("workbook")
    if not isinstance(workbook, dict):
        return contradictions

    funding = workbook.get("funding")
    if not isinstance(funding, dict):
        funding = {}

    def _num(v):
        return v if isinstance(v, (int, float)) and not isinstance(v, bool) else None

    def _fig(key):
        fig = funding.get(key)
        if isinstance(fig, dict):
            return _num(fig.get("value"))
        return _num(fig)

    # --- CHECK 1: funding reconciliation (Section IV request vs fund-table total) ---
    tcep_req = _fig("tcep_request")
    fund_total = _num(funding.get("fund_table_tcep_sccp_total"))
    if tcep_req is not None and fund_total is not None:
        gap = tcep_req - fund_total
        if abs(gap) > 1e-6:
            contradictions.append({
                "field": "TCEP funding request",
                "detail": (
                    f"Section IV TCEP request (${int(tcep_req):,} thousand) does not reconcile with "
                    f"the fund-table TCEP/SCCP total (${int(fund_total):,} thousand); gap of "
                    f"${abs(int(gap)):,} thousand."
                ),
                "values": {
                    "section_IV_request_1000s": int(tcep_req),
                    "fund_table_total_1000s": int(fund_total),
                    "gap_1000s": int(gap),
                },
                "severity": "high",
            })

    # --- CHECK 2: units anomaly (a $1,000s figure that is implausibly small) ---
    figure_keys = ["tcep_request", "tcep_regional_request", "tcep_state_request",
                   "sccp_request", "uncommitted"]
    fig_vals = {k: _fig(k) for k in figure_keys}
    siblings = [v for v in fig_vals.values() if v is not None]
    max_sibling = max(siblings) if siblings else None
    if max_sibling is not None and max_sibling >= 1000:
        for k, v in fig_vals.items():
            if v is None or v <= 0 or v >= 100:
                continue
            fig = funding.get(k)
            cell = fig.get("cell") if isinstance(fig, dict) else None
            contradictions.append({
                "field": f"{k} (units)",
                "detail": (
                    f"'{k}' value {v} is implausibly small next to sibling funding figures in the "
                    f"thousands (max {int(max_sibling):,}); likely a units error ($ vs $1,000s)."
                ),
                "values": {"value": v, "max_sibling_1000s": int(max_sibling), "cell": cell},
                "severity": "medium",
            })

    # --- CHECK 3: impossible milestone dates (template/closeout artifacts) ---
    answers = workbook.get("answers")
    if isinstance(answers, dict):
        seen = set()
        date_re = re.compile(
            r"\b\d{1,2}[/-]\d{1,2}[/-](\d{4})\b|\b(\d{4})[/-]\d{1,2}[/-]\d{1,2}\b"
        )
        for key, entry in answers.items():
            if not isinstance(entry, dict):
                continue
            ans = entry.get("answer")
            years = []
            if isinstance(ans, (datetime.datetime, datetime.date)):
                years.append(ans.year)
                token = ans.isoformat()[:10]
            else:
                token = str(ans)
                for m in date_re.finditer(token):
                    y = m.group(1) or m.group(2)
                    if y:
                        years.append(int(y))
            for y in years:
                if y < 2000 or y > 2040:
                    dedup = (y, token)
                    if dedup in seen:
                        continue
                    seen.add(dedup)
                    contradictions.append({
                        "field": "milestone date",
                        "detail": (
                            f"Answer for line {key} contains an out-of-range date year {y} "
                            f"('{token}'); likely a template/closeout artifact."
                        ),
                        "values": {"line": str(key), "year": y, "raw": token},
                        "severity": "low",
                    })

    # --- CHECK 3b: impossible dates on the funding-info sheet (where the real samples hide
    #     their 1930s closeout artifacts, e.g. 'Project Funding Info #1'!V33) ---
    seen_ms = set()
    for ms in workbook.get("funding_milestones") or []:
        if not isinstance(ms, dict):
            continue
        try:
            y = int(ms.get("year"))
        except (TypeError, ValueError):
            continue
        if y < 2000 or y > 2040:
            token = str(ms.get("value"))
            if (y, token) in seen_ms:
                continue
            seen_ms.add((y, token))
            contradictions.append({
                "field": "milestone date",
                "detail": (
                    f"Funding-sheet cell {ms.get('cell')} holds an out-of-range date year {y} "
                    f"('{token}'); likely a template/closeout artifact."
                ),
                "values": {"cell": str(ms.get("cell")), "year": y, "raw": token},
                "severity": "low",
            })

    return contradictions


def flag_vague_language(narratives) -> list:
    """Flag non-specific / unsupported narrative spans (spec §8/E). Adapts the ai_content_detector
    "generic where specifics should exist" pattern.

    Returns: list of {"source", "span", "reason"}.
    """
    try:
        # --- Normalize input to a list of {"name", "text"} ---
        sources = []
        if narratives is None:
            return []
        if isinstance(narratives, str):
            sources = [{"name": "narrative", "text": narratives}]
        elif isinstance(narratives, dict):
            sources = [{"name": str(narratives.get("name", "narrative")),
                        "text": str(narratives.get("text", ""))}]
        elif isinstance(narratives, (list, tuple)):
            for i, item in enumerate(narratives):
                if isinstance(item, str):
                    sources.append({"name": f"narrative_{i + 1}", "text": item})
                elif isinstance(item, dict):
                    sources.append({"name": str(item.get("name", f"narrative_{i + 1}")),
                                    "text": str(item.get("text", ""))})
        else:
            return []

        sources = [s for s in sources if s.get("text", "").strip()]
        if not sources:
            return []

        # Truncate each source to ~6000 chars to bound the prompt.
        parts = []
        for s in sources:
            parts.append(f"### SOURCE: {s['name']}\n{s['text'][:6000]}")
        combined = "\n\n".join(parts)

        system_prompt = (
            "You are a grant-application reviewer for Caltrans SB1 TCEP/SCCP nominations. "
            "Identify spans of narrative text that are vague, non-specific, or unsupported where "
            "specifics should exist: generic boilerplate where concrete detail is expected, and "
            "quantitative claims (numbers, percentages, benefit figures) stated without a cited "
            "source or backing document. Do not flag well-supported or appropriately general text. "
            'Return ONLY JSON of the form: {"flags": [{"source": "<source name>", '
            '"span": "<the exact vague/unsupported text>", "reason": "<why it is flagged>"}]}. '
            'If nothing warrants flagging, return {"flags": []}.'
        )
        user_prompt = (
            "Review the following narrative source(s) and return the JSON described.\n\n"
            f"{combined}"
        )

        client = _get_client("gpt-4o")
        resp = client.chat.completions.create(
            model="gpt-4o",
            temperature=0.0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        content = resp.choices[0].message.content
        data = _extract_json(content)
        raw_flags = data.get("flags", []) if isinstance(data, dict) else []

        out = []
        for fl in raw_flags:
            if not isinstance(fl, dict):
                continue
            out.append({
                "source": str(fl.get("source", "")),
                "span": str(fl.get("span", "")),
                "reason": str(fl.get("reason", "")),
            })
        return out
    except Exception:  # noqa: BLE001 — offline/CI must never break on this optional LLM pass
        return []


# ==============================================================================
# REPORTING
# ==============================================================================
def build_program_fit_excel(evaluation: dict, rating: dict, eligibility: dict,
                            contradictions: list, project_name: str) -> BytesIO:
    """Build the reviewable Excel report. Reuses the openpyxl style bank _get_styles() and the
    build_evaluation_excel_v2 structure from PDE: a Summary sheet + a per-criterion detail sheet.

    Returns: BytesIO of the .xlsx.
    """
    from openpyxl.styles import Alignment, Font
    from src.project_delivery_evaluator import _get_styles

    s = _get_styles()
    hdr_fill = s['hdr_fill']
    hdr_font = s['hdr_font']
    even_fill = s['even_fill']
    bdr = s['bdr']
    wrap = s['wrap']
    center = s['center']
    top_left = s['top_left']
    bold = s['bold']
    title_font = Font(bold=True, size=14, color="1F4E79")

    evaluation = evaluation if isinstance(evaluation, dict) else {}
    rating = rating if isinstance(rating, dict) else {}
    eligibility = eligibility if isinstance(eligibility, dict) else {}
    contradictions = contradictions if isinstance(contradictions, list) else []
    project_name = str(project_name) if project_name else ""

    def _s(v):
        return "" if v is None else str(v)

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # =====================================================================
    # Sheet 1 — Program Fit Summary
    # =====================================================================
    ws = wb.create_sheet("Program Fit Summary")
    ws.column_dimensions['A'].width = 24
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 40

    row = 1
    # Title
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    tc = ws.cell(row=row, column=1,
                 value=f"Program Fit Evaluation — {project_name or _s(evaluation.get('project_name'))}")
    tc.font = title_font
    tc.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 30
    row += 2

    def _kv(label, value):
        nonlocal row
        lc = ws.cell(row=row, column=1, value=label)
        lc.font = bold
        lc.border = bdr
        lc.alignment = top_left
        vc = ws.cell(row=row, column=2, value=_s(value))
        vc.border = bdr
        vc.alignment = wrap
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
        for c in range(3, 5):
            ws.cell(row=row, column=c).border = bdr
        row += 1

    def _section(title):
        nonlocal row
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        hc = ws.cell(row=row, column=1, value=title)
        hc.font = hdr_font
        hc.fill = hdr_fill
        hc.alignment = center
        for c in range(1, 5):
            ws.cell(row=row, column=c).fill = hdr_fill
            ws.cell(row=row, column=c).border = bdr
        row += 1

    # Project identity block
    _section("Project")
    _kv("Project Name", project_name or evaluation.get("project_name"))
    _kv("District", evaluation.get("district"))
    _kv("Program", evaluation.get("program"))
    row += 1

    # Overall rating block
    _section("Overall Rating")
    _kv("Rating", rating.get("rating"))
    avg = rating.get("average")
    try:
        avg_disp = f"{float(avg):.2f}"
    except (TypeError, ValueError):
        avg_disp = _s(avg)
    _kv("Average Score (1-5)", avg_disp)
    row += 1

    # Eligibility block
    _section("Eligibility")
    _kv("Eligible", eligibility.get("eligible"))
    _kv("Funding Risk", eligibility.get("funding_risk"))
    _kv("Schedule Risk", eligibility.get("schedule_risk"))

    issues = []
    for fail in eligibility.get("failures") or []:
        if isinstance(fail, dict):
            issues.append(f"FAILURE: {_s(fail.get('rule'))} — {_s(fail.get('detail'))}")
        else:
            issues.append(f"FAILURE: {_s(fail)}")
    for warn in eligibility.get("warnings") or []:
        issues.append(f"WARNING: {_s(warn)}")
    if issues:
        for issue in issues:
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            ic = ws.cell(row=row, column=1, value=issue)
            ic.border = bdr
            ic.alignment = wrap
            for c in range(2, 5):
                ws.cell(row=row, column=c).border = bdr
            ws.row_dimensions[row].height = 30
            row += 1
    else:
        _kv("Failures / Warnings", "None")
    row += 1

    # Contradictions block
    _section("Contradictions")
    if contradictions:
        # small table header
        for ci, h in enumerate(("Field", "Severity", "Detail"), 1):
            c = ws.cell(row=row, column=ci, value=h)
            c.font, c.fill, c.border, c.alignment = hdr_font, hdr_fill, bdr, center
        ws.cell(row=row, column=4).border = bdr
        row += 1
        for i, con in enumerate(contradictions):
            con = con if isinstance(con, dict) else {}
            fill = even_fill if row % 2 == 0 else None
            vals = [_s(con.get("field")), _s(con.get("severity")), _s(con.get("detail"))]
            for ci, v in enumerate(vals, 1):
                c = ws.cell(row=row, column=ci, value=v)
                c.border, c.alignment = bdr, wrap
                if fill:
                    c.fill = fill
            ws.cell(row=row, column=4).border = bdr
            if fill:
                ws.cell(row=row, column=4).fill = fill
            row += 1
    else:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        nc = ws.cell(row=row, column=1, value="None detected")
        nc.border, nc.alignment = bdr, wrap
        for c in range(2, 5):
            ws.cell(row=row, column=c).border = bdr
        row += 1

    # =====================================================================
    # Sheet 2 — Criteria Detail
    # =====================================================================
    ws2 = wb.create_sheet("Criteria Detail")
    headers = ["Group", "Criterion", "Score (1-5)", "Confidence",
               "Missing Info", "Evidence / Reasoning"]
    widths = [22, 32, 12, 12, 12, 70]
    for ci, w in enumerate(widths, 1):
        ws2.column_dimensions[openpyxl.utils.get_column_letter(ci)].width = w

    for ci, h in enumerate(headers, 1):
        c = ws2.cell(row=1, column=ci, value=h)
        c.font, c.fill, c.border, c.alignment = hdr_font, hdr_fill, bdr, center
    ws2.freeze_panes = "A2"

    r = 2
    for crit in evaluation.get("criteria") or []:
        crit = crit if isinstance(crit, dict) else {}
        cid = crit.get("criterion_id", crit.get("id", ""))
        name = crit.get("name", "")
        criterion_label = f"{_s(cid)} — {_s(name)}".strip(" —") if cid or name else ""
        conf = crit.get("confidence")
        try:
            conf_disp = f"{float(conf):.2f}"
        except (TypeError, ValueError):
            conf_disp = _s(conf)
        missing = crit.get("missing_info")
        missing_disp = "Yes" if missing else ("" if missing is None else "No")
        vals = [
            _s(crit.get("group")),
            criterion_label,
            _s(crit.get("score")),
            conf_disp,
            missing_disp,
            _s(crit.get("source_reasoning")),
        ]
        fill = even_fill if r % 2 == 0 else None
        for ci, v in enumerate(vals, 1):
            c = ws2.cell(row=r, column=ci, value=v)
            c.border = bdr
            c.alignment = wrap if ci == 6 else center
            if fill:
                c.fill = fill
        r += 1

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
